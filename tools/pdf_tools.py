import PyPDF2
from fpdf import FPDF
from docx import Document

def merge_pdfs(files, output_path):
    merger = PyPDF2.PdfMerger()
    for f in files:
        merger.append(f)
    merger.write(output_path)
    merger.close()

def split_pdf(file, output_folder):
    reader = PyPDF2.PdfReader(file)
    output_files = []
    for i, page in enumerate(reader.pages):
        writer = PyPDF2.PdfWriter()
        writer.add_page(page)
        output_file = f"{output_folder}/page_{i+1}.pdf"
        with open(output_file, "wb") as out:
            writer.write(out)
        output_files.append(output_file)
    return output_files

def word_to_pdf(file, output_path):
    doc = Document(file)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)
    pdf.output(output_path)

def pdf_to_word(file, output_path):
    reader = PyPDF2.PdfReader(file)
    doc = Document()
    for page in reader.pages:
        text = page.extract_text()
        doc.add_paragraph(text)
    doc.save(output_path)
