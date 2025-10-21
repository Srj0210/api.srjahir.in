import os
import re
import subprocess
import threading
from flask import Flask, request, send_file, jsonify, after_this_request
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from fpdf import FPDF

app = Flask(__name__)
CORS(app)

# ✅ Use /tmp for Render-safe read/write operations
UPLOAD_FOLDER = "/tmp/uploads"
OUTPUT_FOLDER = "/tmp/outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def safe_remove(path):
    """Safely remove file if exists"""
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception as e:
        print(f"⚠️ Failed to delete {path}: {e}")


def cleanup_files(paths, delay=5):
    """Delete multiple files after small delay"""
    import time
    time.sleep(delay)
    for p in paths:
        safe_remove(p)


def parse_pages_spec(spec: str, total_pages: int):
    """Parse user-supplied pages spec like 1,3,5-7 or 5689"""
    if not spec:
        return []
    spec = spec.strip()
    pages = set()

    if "," in spec or "-" in spec:
        parts = [p.strip() for p in spec.split(",") if p.strip()]
        for part in parts:
            if "-" in part:
                try:
                    a, b = part.split("-", 1)
                    a, b = int(a), int(b)
                    if a > b:
                        a, b = b, a
                    for n in range(a, b + 1):
                        pages.add(n)
                except:
                    continue
            else:
                try:
                    pages.add(int(part))
                except:
                    continue
    else:
        if spec.isdigit():
            if len(spec) <= 4:
                try:
                    full = int(spec)
                    if 1 <= full <= total_pages:
                        pages.add(full)
                    else:
                        for ch in spec:
                            pages.add(int(ch))
                except:
                    for ch in spec:
                        pages.add(int(ch))
            else:
                for ch in spec:
                    pages.add(int(ch))
        else:
            found = re.findall(r"\d+", spec)
            for f in found:
                pages.add(int(f))

    return sorted([p for p in pages if 1 <= p <= total_pages])


# ---------------------------
# Word → PDF (LibreOffice)
# ---------------------------
@app.route("/word-to-pdf", methods=["POST"])
def word_to_pdf():
    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        filename = secure_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)

        output_filename = os.path.splitext(filename)[0] + ".pdf"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", OUTPUT_FOLDER, input_path
        ], check=True)

        @after_this_request
        def cleanup(response):
            safe_remove(input_path)
            safe_remove(output_path)
            return response

        return send_file(output_path, as_attachment=True, download_name=output_filename)

    except subprocess.CalledProcessError:
        safe_remove(input_path)
        return jsonify({"error": "Conversion failed — LibreOffice error"}), 500
    except Exception as e:
        safe_remove(input_path)
        return jsonify({"error": str(e)}), 500


# ---------------------------
# PDF → Word (Hybrid Pro Engine v2, Fixed)
# ---------------------------
@app.route("/pdf-to-word", methods=["POST"])
def pdf_to_word():
    import fitz  # PyMuPDF
    from pdf2docx import Converter
    from pdf2image import convert_from_path
    import pytesseract
    from docx import Document

    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        filename = secure_filename(file.filename)
        input_path = os.path.join("/tmp", filename)
        file.save(input_path)

        output_name = os.path.splitext(filename)[0] + ".docx"
        output_path = os.path.join("/tmp", output_name)

        # STEP 1️⃣: Detect if PDF contains text (Safe version)
        text_based = False
        try:
            doc = fitz.open(input_path)
            for page in doc:
                extracted = page.get_text("text")
                if extracted and extracted.strip():
                    text_based = True
                    break
            doc.close()
        except Exception as e:
            print("⚠️ Text detection failed:", e)

        # STEP 2️⃣: Use pdf2docx for text PDFs
        if text_based:
            print("🧩 Text-based PDF detected — using pdf2docx engine.")
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()

        # STEP 3️⃣: Fallback OCR for image PDFs
        else:
            print("🧠 Image-based PDF detected — using OCR fallback.")
            images = convert_from_path(input_path, dpi=200)
            document = Document()

            for page_num, img in enumerate(images, start=1):
                text = pytesseract.image_to_string(img)
                document.add_paragraph(f"--- Page {page_num} ---")
                document.add_paragraph(text.strip() if text.strip() else "[No text found]")
                document.add_page_break()

            document.save(output_path)

        # STEP 4️⃣: Verify output
        if not os.path.exists(output_path) or os.path.getsize(output_path) < 2000:
            raise Exception("Output file empty or corrupted.")

        os.chmod(output_path, 0o777)

        @after_this_request
        def cleanup(response):
            safe_remove(input_path)
            safe_remove(output_path)
            return response

        return send_file(output_path, as_attachment=True, download_name=output_name)

    except Exception as e:
        print(f"❌ PDF→Word Hybrid Error: {e}")
        return jsonify({"error": str(e)}), 500


# ---------------------------
# Split PDF (custom pages)
# ---------------------------
@app.route("/split-pdf", methods=["POST"])
def split_pdf():
    try:
        file = request.files.get("file")
        pages_spec = request.form.get("pages", "").strip()
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        filename = secure_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        if not pages_spec:
            selected = list(range(1, total_pages + 1))
        else:
            selected = parse_pages_spec(pages_spec, total_pages)

        if not selected:
            safe_remove(input_path)
            return jsonify({"error": "No valid pages selected"}), 400

        writer = PdfWriter()
        for p in selected:
            idx = p - 1
            if 0 <= idx < total_pages:
                writer.add_page(reader.pages[idx])

        output_filename = f"split_{'_'.join(str(x) for x in selected)}.pdf"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        with open(output_path, "wb") as f_out:
            writer.write(f_out)

        @after_this_request
        def cleanup(response):
            safe_remove(input_path)
            safe_remove(output_path)
            return response

        return send_file(output_path, as_attachment=True, download_name=output_filename)

    except Exception as e:
        safe_remove(input_path)
        return jsonify({"error": str(e)}), 500


# ---------------------------
# Text → PDF
# ---------------------------
@app.route("/text-to-pdf", methods=["POST"])
def text_to_pdf():
    try:
        text = request.form.get("text", "")
        if not text.strip():
            return jsonify({"error": "No text provided"}), 400

        output_filename = f"text_{os.urandom(6).hex()}.pdf"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=12)
        for line in text.splitlines():
            pdf.multi_cell(0, 10, line)
        pdf.output(output_path)

        @after_this_request
        def cleanup(response):
            safe_remove(output_path)
            return response

        return send_file(output_path, as_attachment=True, download_name="text.pdf")

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ---------------------------
# Root Endpoint
# ---------------------------
@app.route("/", methods=["GET"])
def home():
    return jsonify({"status": "SRJahir Tools API", "ready": True})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 10000)))